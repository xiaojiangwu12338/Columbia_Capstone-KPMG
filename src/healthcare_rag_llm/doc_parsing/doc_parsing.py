import os
import json
import datetime
from pathlib import Path
import csv  # NEW

from pdfminer.high_level import extract_text as pdfminer_extract_text
from docx import Document
import pdfplumber

try:
    import pytesseract
    from PIL import Image
    from PIL import ImageOps, ImageEnhance
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False


# Common watermark keywords
WATERMARK_HINTS = {"ARCHIVE", "DRAFT", "CONFIDENTIAL", "SAMPLE"}
WATERMARK_LETTERS = {ch.lower() for wm in WATERMARK_HINTS for ch in wm.lower()}


def pdfminer_pages_text(path, n_pages=None):
    """Extract text per page using pdfminer"""
    pages = []
    if n_pages is None:
        with pdfplumber.open(str(path)) as pdf:
            n_pages = len(pdf.pages)
    for i in range(n_pages):
        t = pdfminer_extract_text(str(path), page_numbers=[i]) or ""
        pages.append(t)
    return pages


def detect_watermarks_per_page(pages_text, hints=WATERMARK_HINTS):
    """
    Detect watermarks from pdfminer plain text:
    - Only trigger if keyword appears as a standalone line
    - Or appears split across multiple consecutive lines
    """
    results = []

    for i, page_text in enumerate(pages_text, start=1):
        lines = [l.strip() for l in page_text.splitlines() if l.strip()]
        lowered = [l.lower() for l in lines]

        for h in hints:
            h_low = h.lower()

            # Case 1: full keyword appears as its own line
            if any(l == h_low for l in lowered):
                results.append({"text": h.upper(), "pages": [i]})
                continue

            # Case 2: keyword spelled out letter by letter on consecutive lines
            letters = list(h_low)
            for j in range(len(lowered) - len(letters) + 1):
                segment = lowered[j : j + len(letters)]
                if segment == letters:
                    results.append({"text": h.upper(), "pages": [i]})
                    break

    return results


def clean_page_text_remove_isolated_letters(page_text):
    """Remove isolated watermark letters/words in text"""
    cleaned_lines = []
    for line in page_text.splitlines():
        stripped = line.strip().lower()
        if len(stripped) == 1 and stripped in WATERMARK_LETTERS:
            continue
        cleaned_lines.append(line)
    return "\n".join(cleaned_lines)


# NEW: lightweight category lookup (no pandas)
def _lookup_category(file_name, csv_path="data/metadata/metadata_filled.csv"):
    """
    Return category for a given file_name by scanning a CSV that has columns
    'file_name' and 'category'. Header and values are whitespace-trimmed.
    If CSV is missing or no match is found, return 'unknown'.
    """
    try:
        with open(csv_path, "r", encoding="utf-8", newline="") as f:
            reader = csv.DictReader(f)
            # Normalize header names by stripping whitespace
            if reader.fieldnames:
                reader.fieldnames = [h.strip() if h else h for h in reader.fieldnames]
            for raw in reader:
                # Normalize keys & values
                row = { (k.strip() if k else k): (v.strip() if isinstance(v, str) else v)
                        for k, v in raw.items() }
                # Handle possible trailing space in 'file_name ' column name
                fn = row.get("file_name")
                if fn is None:
                    fn = row.get("file_name") or row.get("file_name_") or row.get("file") or row.get("filename")
                if fn == file_name:
                    cat = row.get("category") or row.get("Category") or row.get("CATEGORY")
                    return cat if (isinstance(cat, str) and cat.strip()) else "unknown"
    except Exception:
        pass
    return "unknown"


def parse_file(file_path, save_txt=True, save_json=True, out_dir="docs", metadata_csv_path="data/metadata/metadata_filled.csv"):
    """
    Parse a document into text + metadata + tables + watermarks + OCR fallback.
    OCR fallback: only triggered when the page contains images, runs per image,
    results stored separately without contaminating main text.
    """
    file_path = Path(file_path)
    ext = file_path.suffix.lower()



    # Ensure output directory exists

    Path(out_dir).mkdir(parents=True, exist_ok=True)



    base_name = file_path.stem

    json_path = Path(out_dir) / f"{base_name}.json"



    # ---------- NEW SKIP LOGIC WITH CONSOLE OUTPUT ----------
    if json_path.exists():
        try:
            with open(json_path, "r", encoding="utf-8") as f:
                existing_metadata = json.load(f)
            print(f"[SKIPPED] {file_path.name} already parsed → using existing JSON.")
            return existing_metadata
        except Exception:
            print(f"[WARN] Existing JSON for {file_path.name} is unreadable → re-parsing.")
            # Continue to normal parse
    # ---------------------------------------------------------


    
    all_text = []
    # NEW: determine category (best-effort)
    category = _lookup_category(file_path.name, metadata_csv_path)

    metadata = {
        "source_file": str(file_path),
        "file_name": file_path.name,
        "extension": ext,
        "parsed_time": datetime.datetime.now().isoformat(timespec="seconds"),
        "pages": [],
        "tables": [],
        "watermarks": [],
        "ocr": False,
        "category": category,  # NEW
    }

    try:
        if ext == ".pdf":
            plumber_pages = []
            with pdfplumber.open(str(file_path)) as pdf:
                total_pages = len(pdf.pages)
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text() or ""
                    plumber_pages.append(page_text)

            # Watermark detection
            pdfminer_pages = pdfminer_pages_text(str(file_path), n_pages=total_pages)
            watermarks = detect_watermarks_per_page(pdfminer_pages, WATERMARK_HINTS)
            metadata["watermarks"] = watermarks

            # Generate clean text & OCR fallback
            with pdfplumber.open(str(file_path)) as pdf:
                for idx, page in enumerate(pdf.pages):
                    page_no = idx + 1
                    raw_page_text = plumber_pages[idx] or pdfminer_pages[idx] or ""
                    if not raw_page_text.strip() and not page.images:
                        continue

                    # Remove isolated watermark letters/words
                    clean_page_text = clean_page_text_remove_isolated_letters(raw_page_text)

                    # Table extraction
                    try:
                        tables = page.extract_tables()
                        for tbl in tables:
                            metadata["tables"].append({"page": page_no, "table": tbl})
                    except Exception:
                        pass

                    page_record = {
                        "page": page_no,
                        "text": clean_page_text,
                        "ocr_fallback": []
                    }

                    # OCR fallback: run per image
                    if OCR_AVAILABLE and page.images:
                        for img_obj in page.images:
                            try:
                                x0, top, x1, bottom = img_obj["x0"], img_obj["top"], img_obj["x1"], img_obj["bottom"]
                                cropped = page.within_bbox((x0, top, x1, bottom)).to_image(resolution=300)
                                pil_img = cropped.original.convert("L")
                                pil_img = ImageOps.invert(pil_img)
                                pil_img = ImageEnhance.Contrast(pil_img).enhance(2.0)
                                text_img = pytesseract.image_to_string(pil_img)
                                if text_img.strip():
                                    page_record["ocr_fallback"].append({
                                        "bbox": [x0, top, x1, bottom],
                                        "text": text_img.strip()
                                    })
                            except Exception:
                                continue

                    if clean_page_text.strip() or page_record["ocr_fallback"]:
                        metadata["pages"].append(page_record)
                        all_text.append(clean_page_text)

        elif ext == ".docx":
            doc = Document(str(file_path))
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    metadata["pages"].append({"page": i+1, "text": para.text, "ocr_fallback": []})
                    all_text.append(para.text)
            for t in doc.tables:
                table_data = []
                for row in t.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                if table_data:
                    metadata["tables"].append({"page": None, "table": table_data})

        elif ext == ".doc":
            raise ValueError("DOC format not supported. Please convert to DOCX first.")

        elif ext == ".txt":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                lines = f.readlines()
            for i, line in enumerate(lines):
                if line.strip():
                    metadata["pages"].append({"page": i+1, "text": line.strip(), "ocr_fallback": []})
                    all_text.append(line.strip())

        else:
            raise ValueError(f"Unsupported file type: {ext}")

        metadata["full_text"] = "\n\n".join(all_text)

    except Exception as e:
        metadata["error"] = str(e)

    metadata["ocr"] = any(
        bool(p.get("ocr_fallback")) for p in metadata["pages"]
    )

    # Save results
    Path(out_dir).mkdir(parents=True, exist_ok=True)
    base_name = file_path.stem

    if save_txt:
        txt_path = Path(out_dir) / f"{base_name}.txt"
        with open(txt_path, "w", encoding="utf-8", errors="ignore") as f:
            f.write(metadata.get("full_text", ""))

    if save_json:
        json_path = Path(out_dir) / f"{base_name}.json"
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(metadata, f, ensure_ascii=False, indent=2)

    return metadata


# Example
if __name__ == "__main__":
    input_pdf = "data/raw/Childrens Evolution of Care/State/CFTSS/service_delivery_designation_guidance_cftss_hcbs.pdf"
    result = parse_file(input_pdf, save_txt=True, save_json=True)

    print("Parsed file:", result["file_name"])
    print("Category:", result.get("category"))  # NEW
    print("Pages:", len(result["pages"]))
    print("Tables:", len(result["tables"]))
    print("Watermarks:", result["watermarks"])
    print("OCR:", result["ocr"])
    print("Total chars:", len(result.get("full_text", "")))
