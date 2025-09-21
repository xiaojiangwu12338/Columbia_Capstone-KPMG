import os
import json
import datetime
from pathlib import Path

from pdfminer.high_level import extract_text as pdfminer_extract_text
from docx import Document
import pdfplumber

try:
    import pytesseract
    from PIL import Image
    from pdf2image import convert_from_path
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False


# common watermarks
WATERMARK_HINTS = {"ARCHIVE", "DRAFT", "CONFIDENTIAL", "SAMPLE"}
WATERMARK_LETTERS = {ch.lower() for wm in WATERMARK_HINTS for ch in wm.lower()}


def pdfminer_pages_text(path, n_pages=None):
    """ use pdfminer to extract text by page"""
    pages = []
    if n_pages is None:
        with pdfplumber.open(str(path)) as pdf:
            n_pages = len(pdf.pages)
    for i in range(n_pages):
        t = pdfminer_extract_text(str(path), page_numbers=[i]) or ""
        pages.append(t)
    return pages


def detect_watermarks_per_page(pages_text, hints=WATERMARK_HINTS):
    """
    use pdfminer to check the watermarks
    return watermark and page
    sample return: [{"text": "ARCHIVE", "pages": [1,2]}]
    """
    found = {}
    for i, page_text in enumerate(pages_text, start=1):
        low = page_text.lower()
        for h in hints:
            if h.lower() in low:
                found.setdefault(h.upper(), []).append(i)
    return [{"text": k, "pages": v} for k, v in found.items()]


def clean_page_text_remove_isolated_letters(page_text):
    """
    delete isolated letters for watermarks
    """
    cleaned_lines = []
    for line in page_text.splitlines():
        stripped = line.strip().lower()
        if len(stripped) == 1 and stripped in WATERMARK_LETTERS:
            continue
        cleaned_lines.append(line)
    return "\n".join(cleaned_lines)


def parse_file(file_path, save_txt=True, save_json=True, out_dir="docs"):
    """
    Parse a document (PDF/DOCX/TXT) into text + metadata + tables + watermarks.
    """
    file_path = Path(file_path)
    ext = file_path.suffix.lower()
    all_text = []
    metadata = {
        "source_file": str(file_path),
        "file_name": file_path.name,
        "extension": ext,
        "parsed_time": datetime.datetime.now().isoformat(timespec="seconds"),
        "pages": [],
        "tables": [],
        "watermarks": []
    }

    try:
        if ext == ".pdf":
            plumber_pages = []
            with pdfplumber.open(str(file_path)) as pdf:
                total_pages = len(pdf.pages)
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text() or ""
                    plumber_pages.append(page_text)
                    # extract tables
                    try:
                        tables = page.extract_tables()
                        for tbl in tables:
                            metadata["tables"].append({"page": i+1, "table": tbl})
                    except Exception:
                        pass

            # use pdfminer to check watermarks
            pdfminer_pages = pdfminer_pages_text(str(file_path), n_pages=total_pages)
            watermarks = detect_watermarks_per_page(pdfminer_pages, WATERMARK_HINTS)
            metadata["watermarks"] = watermarks

            # generate clean text without watermarks
            for idx in range(total_pages):
                page_no = idx + 1
                raw_page_text = plumber_pages[idx] or pdfminer_pages[idx] or ""
                if not raw_page_text.strip():
                    continue
                clean_page_text = clean_page_text_remove_isolated_letters(raw_page_text)
                if clean_page_text.strip():
                    metadata["pages"].append({"page": page_no, "text": clean_page_text})
                    all_text.append(clean_page_text)

        elif ext == ".docx":
            doc = Document(str(file_path))
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    metadata["pages"].append({"page": i+1, "text": para.text})
                    all_text.append(para.text)
            for t in doc.tables:
                table_data = []
                for row in t.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                if table_data:
                    metadata["tables"].append({"page": None, "table": table_data})

        elif ext == ".doc":
            raise ValueError("DOC format not supported. Please convert to DOCX first.")

        elif ext == ".txt":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                lines = f.readlines()
            for i, line in enumerate(lines):
                if line.strip():
                    metadata["pages"].append({"page": i+1, "text": line.strip()})
                    all_text.append(line.strip())

        else:
            raise ValueError(f"Unsupported file type: {ext}")

        metadata["full_text"] = "\n\n".join(all_text)

    except Exception as e:
        metadata["error"] = str(e)

    # save result
    Path(out_dir).mkdir(parents=True, exist_ok=True)
    base_name = file_path.stem

    if save_txt:
        txt_path = Path(out_dir) / f"{base_name}.txt"
        with open(txt_path, "w", encoding="utf-8", errors="ignore") as f:
            f.write(metadata.get("full_text", ""))

    if save_json:
        json_path = Path(out_dir) / f"{base_name}.json"
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(metadata, f, ensure_ascii=False, indent=2)

    return metadata

# use case 
if __name__ == "__main__":
    input_pdf = "data/Childrens Evolution of Care/State/CFTSS/2023-09-05_bill_req_update_child_CFTSS.pdf"
    result = parse_file(input_pdf, save_txt=True, save_json=True)
    print("Parsed file:", result["file_name"])
    print("Pages:", len(result["pages"]))
    print("Tables:", len(result["tables"]))
    print("Watermarks:", result["watermarks"])
